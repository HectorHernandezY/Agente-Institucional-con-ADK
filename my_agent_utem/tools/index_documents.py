from __future__ import annotations
import os
from typing import Any, Optional
import hashlib
from datetime import datetime
import io

from google.cloud import firestore
from google.cloud import storage
from vertexai.language_models import TextEmbeddingModel
from google.adk.tools import FunctionTool
from google.auth import default

# Para leer DOCX
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx no instalado. Instala con: pip install python-docx")

# Configuración
PROJECT_ID = os.getenv("FIRESTORE_PROJECT_ID", "muruna-utem-project")
GCS_BUCKET = os.getenv("GCS_RAG_BUCKET", "db_agent_utem")
VERTEX_LOCATION = os.getenv("GOOGLE_CLOUD_LOCATION", "us-central1")

# Configurar credenciales con quota project
credentials, _ = default(quota_project_id=PROJECT_ID)

# Inicializar clientes
db = firestore.Client(project=PROJECT_ID, credentials=credentials)
storage_client = storage.Client(project=PROJECT_ID, credentials=credentials)

# Inicializar Vertex AI
import vertexai
vertexai.init(project=PROJECT_ID, location=VERTEX_LOCATION, credentials=credentials)

# ✅ LAZY LOADING del modelo de embeddings
_embedding_model = None

def get_embedding_model():
    """Obtiene el modelo de embeddings con lazy loading"""
    global _embedding_model
    if _embedding_model is None:
        try:
            _embedding_model = TextEmbeddingModel.from_pretrained("text-embedding-004")
            print("✅ Modelo de embeddings cargado: text-embedding-004")
        except Exception as e:
            print(f"⚠️  Fallback a textembedding-gecko@003: {e}")
            _embedding_model = TextEmbeddingModel.from_pretrained("textembedding-gecko@003")
    return _embedding_model


class DocumentProcessor:
    """Procesa documentos y genera chunks"""
    
    def __init__(self, chunk_size: int = 800, overlap: int = 150):
        self.chunk_size = chunk_size
        self.overlap = overlap
    
    def chunk_text(self, text: str) -> list[str]:
        """Divide texto en chunks con overlap"""
        if not text or len(text.strip()) == 0:
            return []
        
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = min(start + self.chunk_size, text_length)
            chunk = text[start:end]
            
            # Evitar cortar palabras a mitad
            if end < text_length and text[end] not in [' ', '\n', '.', ',', '!', '?']:
                last_space = chunk.rfind(' ')
                if last_space > self.chunk_size // 2:  # Si hay espacio en la segunda mitad
                    end = start + last_space
                    chunk = text[start:end]
            
            chunk_clean = chunk.strip()
            if chunk_clean:
                chunks.append(chunk_clean)
            
            start = end - self.overlap if end < text_length else end
        
        return chunks


def extract_text_from_docx(blob_bytes: bytes) -> str:
    """
    Extrae texto de un archivo DOCX de manera eficiente
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx no está instalado. Ejecuta: pip install python-docx")
    
    try:
        # Usar BytesIO para evitar escribir a disco
        docx_stream = io.BytesIO(blob_bytes)
        doc = Document(docx_stream)
        
        # Extraer párrafos
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Extraer tablas
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)
        
        # Unir con doble salto de línea
        full_text = '\n\n'.join(paragraphs)
        
        # Limpiar espacios múltiples
        import re
        full_text = re.sub(r'\s+', ' ', full_text)
        full_text = re.sub(r'\n\s*\n', '\n\n', full_text)
        
        return full_text.strip()
        
    except Exception as e:
        raise RuntimeError(f"Error extrayendo texto de DOCX: {e}")


def read_file_from_gcs(blob_name: str) -> str:
    """Lee contenido de archivo desde GCS según su tipo"""
    try:
        bucket = storage_client.bucket(GCS_BUCKET)
        blob = bucket.blob(blob_name)
        
        file_extension = blob_name.lower().split('.')[-1]
        
        # Para archivos DOCX
        if file_extension == 'docx':
            print(f"    📄 Detectado DOCX, extrayendo texto...")
            blob_bytes = blob.download_as_bytes()
            text = extract_text_from_docx(blob_bytes)
            # Liberar memoria del blob
            del blob_bytes
            return text
        
        # Para archivos de texto plano
        elif file_extension in ['txt', 'md', 'csv']:
            try:
                return blob.download_as_text(encoding='utf-8')
            except UnicodeDecodeError:
                return blob.download_as_text(encoding='latin-1')
        
        # Para PDFs (requiere PyPDF2 o similar)
        elif file_extension == 'pdf':
            raise NotImplementedError("Soporte para PDF no implementado aún")
        
        else:
            raise ValueError(f"Tipo de archivo no soportado: {file_extension}")
        
    except Exception as e:
        raise RuntimeError(f"Error leyendo archivo {blob_name}: {e}")


def index_document_from_gcs(
    file_path: str,
    document_name: Optional[str] = None
) -> dict[str, Any]:
    """
    Indexa un documento desde GCS.
    """
    try:
        print(f"  → Leyendo archivo: {file_path}")
        
        # 1. Leer y extraer texto
        content = read_file_from_gcs(file_path)
        
        if not content or len(content.strip()) < 10:
            return {
                "ok": False,
                "status": "Error",
                "message": "El documento está vacío o tiene muy poco contenido"
            }
        
        print(f"    📏 Texto extraído: {len(content):,} caracteres")
        
        # 2. Generar ID único
        doc_name = document_name or file_path.split('/')[-1]
        doc_id = hashlib.md5(f"{doc_name}_{datetime.now().isoformat()}".encode()).hexdigest()
        
        # 3. Dividir en chunks
        print(f"  → Generando chunks...")
        processor = DocumentProcessor(chunk_size=800, overlap=150)
        chunks = processor.chunk_text(content)
        
        # LIBERAR MEMORIA del texto original
        del content
        import gc
        gc.collect()
        
        if not chunks:
            return {
                "ok": False,
                "status": "Error",
                "message": "No se pudieron generar chunks del documento"
            }
        
        print(f"    📦 {len(chunks)} chunks creados")
        
        # 4. Obtener modelo de embeddings (lazy loading)
        embedding_model = get_embedding_model()
        
        # 5. Generar embeddings en BATCHES pequeños
        print(f"  → Generando embeddings en batches...")
        BATCH_SIZE = 5  # Procesar de 5 en 5
        all_embeddings = []
        
        for i in range(0, len(chunks), BATCH_SIZE):
            batch_chunks = chunks[i:i + BATCH_SIZE]
            batch_embeddings = embedding_model.get_embeddings(batch_chunks)
            all_embeddings.extend([emb.values for emb in batch_embeddings])
            
            print(f"    ⚙️  Batch {i//BATCH_SIZE + 1}/{(len(chunks)-1)//BATCH_SIZE + 1}")
            
            # Liberar memoria del batch
            del batch_chunks, batch_embeddings
            gc.collect()
        
        # 6. Guardar en Firestore en BATCHES
        print(f"  → Guardando en Firestore...")
        FIRESTORE_BATCH_SIZE = 500  # Firestore max batch size
        chunk_ids = []
        
        for batch_start in range(0, len(chunks), FIRESTORE_BATCH_SIZE):
            batch = db.batch()
            batch_end = min(batch_start + FIRESTORE_BATCH_SIZE, len(chunks))
            
            for idx in range(batch_start, batch_end):
                chunk_id = f"{doc_id}_chunk_{idx}"
                chunk_ids.append(chunk_id)
                
                doc_ref = db.collection("rag_vectores").document(chunk_id)
                
                chunk_data = {
                    "chunk_id": chunk_id,
                    "doc_id": doc_id,
                    "doc_name": doc_name,
                    "gcs_uri": f"gs://{GCS_BUCKET}/{file_path}",
                    "text": chunks[idx],
                    "embedding": all_embeddings[idx],
                    "chunk_index": idx,
                    "total_chunks": len(chunks),
                    "created_at": firestore.SERVER_TIMESTAMP,
                    "access_count": 0
                }
                
                batch.set(doc_ref, chunk_data)
            
            batch.commit()
            print(f"    💾 Guardados chunks {batch_start}-{batch_end}")
        
        # LIBERAR TODO
        del chunks, all_embeddings
        gc.collect()
        
        print(f"  ✓ Completado: {len(chunk_ids)} chunks guardados")
        
        return {
            "ok": True,
            "status": "Documento indexado exitosamente",
            "doc_id": doc_id,
            "doc_name": doc_name,
            "chunks_created": len(chunk_ids),
            "gcs_uri": f"gs://{GCS_BUCKET}/{file_path}"
        }
        
    except Exception as e:
        print(f"  ✗ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            "ok": False,
            "status": "Error",
            "message": f"Error indexando documento: {str(e)}"
        }


# Exportar como tools
index_document_tool = FunctionTool(index_document_from_gcs)